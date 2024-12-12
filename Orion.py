import os
import json
import pandas as pd
import logging
import tkinter as tk
from pathlib import Path
from sklearn.ensemble import IsolationForest
from docx import Document
from datetime import datetime
from tkinter import filedialog
from tkinter import scrolledtext
from tkinter import messagebox
from sklearn.preprocessing import LabelEncoder

# Function to load logs from a directory
def load_logs(directory):
    logs = []
    for filename in os.listdir(directory):
        filepath = os.path.join(directory, filename)
        if filename.endswith('.csv'):
            try:
                df = pd.read_csv(filepath)
                csv_logs = df.to_dict(orient='records')
                logs.extend(csv_logs)
                logging.info(f"Successfully loaded and converted CSV file: {filename}")
            except Exception as e:
                logging.error(f"Error loading and processing CSV file {filename}: {e}")
    return logs

# Function to perform anomaly detection using Isolation Forest
def detect_anomalies(log_data):
    label_encoder_event_type = LabelEncoder()
    label_encoder_source_ip = LabelEncoder()

    log_data['Event Type'] = label_encoder_event_type.fit_transform(log_data['Event Type'])
    log_data['Source IP'] = label_encoder_source_ip.fit_transform(log_data['Source IP'])

    model = IsolationForest(contamination=0.1)
    log_data['anomaly'] = model.fit_predict(log_data[['Event Type', 'Source IP']])
    
    anomalies = log_data[log_data['anomaly'] == -1]
    return anomalies, label_encoder_event_type, label_encoder_source_ip

# GUI for application
class ThreatHuntingGUI:
    def __init__(self, root):
        self.root = root
        self.root.title("Threat Hunter 'OrioN'")
        frame = tk.Frame(root, bg='black')
        frame.place(relwidth=1, relheight=1)
        self.root.geometry("1920x1080")

        self.log_directory = ""  # Default path

        # GUI elements
        self.create_widgets()

    def create_widgets(self):
        """ Creates GUI elements """
        self.dir_label = tk.Label(self.root, text="No Directory Selected")
        self.dir_label.pack(pady=10)

        self.load_button = tk.Button(self.root, text="Select Log Directory", command=self.load_log_directory)
        self.load_button.pack(pady=10)

        self.run_button = tk.Button(self.root, text="Run OrioN", command=self.run_threat_hunting, bg='red', fg='white')
        self.run_button.pack(pady=10)

        self.results_text = scrolledtext.ScrolledText(self.root, wrap=tk.WORD, width=100, height=40)
        self.results_text.pack(padx=10, pady=10)

        self.past_reports_label = tk.Label(self.root, text="Past Generated Reports")
        self.past_reports_label.pack(pady=5)

        self.past_reports_listbox = tk.Listbox(self.root, width=80, height=5)
        self.past_reports_listbox.pack(padx=10, pady=5)

        self.past_reports_listbox.bind("<Double-1>", self.open_selected_report)

    def open_selected_report(self, event):
        """ Opens the selected past report when clicked """
        selected_report = self.past_reports_listbox.get(tk.ACTIVE)
        if selected_report:
            report_filename = os.path.join(self.log_directory, selected_report)
            self.open_report_in_window(report_filename)

    def load_log_directory(self):
        directory = filedialog.askdirectory()
        if directory:
            self.log_directory = directory
            self.dir_label.config(text=f"Selected Directory: {self.log_directory}")
            self.load_button.config(state=tk.DISABLED)
        else:
            messagebox.showwarning("No Directory", "No directory has been selected. Please select a valid directory.")

    def run_threat_hunting(self):
        """ Handles the process of running threat hunting """
        if not self.log_directory:
            messagebox.showwarning("No Directory", "Please select a directory first.")
            return

        # Load and process logs
        logs = load_logs(self.log_directory)
        if not logs:
            messagebox.showwarning("No Logs", "No logs were found in the selected directory")
            return

        log_data = pd.DataFrame(logs)

        # Check for essential columns
        required_columns = ['Event Name', 'Event Receive Time', 'Source IP', 'Destination IP']
        missing_columns = [col for col in required_columns if col not in log_data.columns]

        if missing_columns:
            messagebox.showwarning("Missing Columns", f"Logs are missing necessary columns: {', '.join(missing_columns)}")
            return

        # Detect Anomalies and get label encoders
        anomalies, label_encoder_event_type, label_encoder_source_ip = detect_anomalies(log_data)

        if anomalies.empty:
            messagebox.showinfo("No Anomalies", "No anomalies detected")
            self.results_text.delete(1.0, tk.END)
            self.results_text.insert(tk.END, "No anomalies detected.")
        else:
            # Display Findings
            self.results_text.delete(1.0, tk.END)
            self.results_text.insert(tk.END, f"Anomalies Detected: {len(anomalies)}\n\n")

            # Display Details
            for _, anomaly in anomalies.iterrows():
                # Use inverse_transform to get the original Event Type and Source IP
                event_type = label_encoder_event_type.inverse_transform([anomaly['Event Type']])[0]
                source_ip = label_encoder_source_ip.inverse_transform([anomaly['Source IP']])[0]
                
                self.results_text.insert(tk.END, f"Anomaly Event Name: {anomaly['Event Name']}\n")
                self.results_text.insert(tk.END, f"Anomaly Event Type: {event_type}\n")
                self.results_text.insert(tk.END, f"Anomaly Event Time Received: {anomaly['Event Receive Time']}\n")
                self.results_text.insert(tk.END, f"Anomaly Event Source IP: {source_ip}\n")
                self.results_text.insert(tk.END, f"Anomaly Event Destination IP: {anomaly['Destination IP']}\n")

            # Generate the report and add it to the list of past reports
            report_filename = f"threat_hunting_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            report_filepath = os.path.join(self.log_directory, report_filename)  # Ensure full file path
            
            # Generate the report and save it
            self.generate_report(anomalies, report_filepath, label_encoder_event_type, label_encoder_source_ip)

            # Add full file path to the past reports listbox (instead of just filename)
            self.past_reports_listbox.insert(tk.END, report_filepath)

            # Display message that the report has been generated
            messagebox.showinfo("Report Generated", f"Report generated: {report_filename}")


    def generate_report(self, anomalies, report_filename, label_encoder_event_type, label_encoder_source_ip):
        """ Generates a Word document report of the findings """
        document = Document()
        document.add_heading('Threat Hunting Report', 0)

        document.add_paragraph(f'Report generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        document.add_paragraph("The following anomalies were detected in the logs:\n")

        for _, anomaly in anomalies.iterrows():
            event_type = label_encoder_event_type.inverse_transform([anomaly['Event Type']])[0]
            source_ip = label_encoder_source_ip.inverse_transform([anomaly['Source IP']])[0]

            document.add_heading(f"Anomaly Event Name: {anomaly['Event Name']}", level=1)
            document.add_paragraph(f"Anomaly Event Type: {event_type}")
            document.add_paragraph(f"Event Time Received: {anomaly['Event Receive Time']}")
            document.add_paragraph(f"Event Source IP: {source_ip}")
            document.add_paragraph(f"Event Destination IP: {anomaly['Destination IP']}")
            document.add_paragraph(f"Reason for Detection: Anomalous behavior detected based on event type and IP.")
            document.add_paragraph("Suggested Actions: Investigate the source of traffic. Block suspicious IP if necessary.\n")
            document.add_paragraph("--------------------------------------------------------------------------------------------")

        document.save(report_filename)

    def open_report_in_window(self, report_filename):
        """ Opens the selected report in a new window """
        try:
            # Check if the report file exists
            if not os.path.exists(report_filename):
                messagebox.showerror("File Not Found", f"The file does not exist at the specified location: {report_filename}")
                return

            from docx import Document

            # Open the docx file using the full file path
            doc = Document(report_filename)
            
            # Create a new window to display the report
            report_window = tk.Toplevel(self.root)
            report_window.title(f"Past Report - {Path(report_filename).name}")
            report_window.geometry("800x600")
            
            # Create a ScrolledText widget to display the content of the Word document
            report_text = scrolledtext.ScrolledText(report_window, wrap=tk.WORD, width=100, height=40)
            report_text.pack(padx=10, pady=10)
            
            # Extract text from the docx and insert it into the ScrolledText widget
            for para in doc.paragraphs:
                report_text.insert(tk.END, para.text + "\n")
            
            # Disable editing of the ScrolledText widget
            report_text.config(state=tk.DISABLED)
            
            # Show the window
            report_window.mainloop()

        except Exception as e:
            messagebox.showerror("Error", f"Failed to open the report: {e}")

# Main function to go through the threat hunting process
if __name__ == "__main__":
    root = tk.Tk()
    app = ThreatHuntingGUI(root)
    root.mainloop()
